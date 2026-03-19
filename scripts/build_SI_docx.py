#!/usr/bin/env python3
"""
Build the complete revised Supplementary Information document for Nature Cities R1.
Generates a DOCX file with all existing + new sections, figures, tables, and references.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# =============================================================================
# PATHS
# =============================================================================
BASE = "/Users/kangninghuang/Library/CloudStorage/GoogleDrive-kh3657@nyu.edu/My Drive/Grants_Fellowship/2024 NYU China Grant/0.CleanProject_GlobalScaling"
REV = os.path.join(BASE, "_Revision1")
FIG_PNG = os.path.join(REV, "figures", "png")
OUT_DIR = os.path.join(REV, "docs")

# =============================================================================
# FIGURE MAP
# =============================================================================
FIGS = {
    'S1': os.path.join(FIG_PNG, 'OrigSI_page9.png'),
    'S2': os.path.join(FIG_PNG, 'OrigSI_page10.png'),
    'S3': os.path.join(FIG_PNG, 'OrigSI_page11.png'),
    'S4': os.path.join(FIG_PNG, 'OrigSI_page12.png'),
    'S5': os.path.join(FIG_PNG, 'OrigSI_page13.png'),
    'S6': os.path.join(FIG_PNG, 'OrigSI_page14.png'),
    'S7': os.path.join(FIG_PNG, 'Fig2_Decentered_Source_Sensitivity_2026-02-11.png'),
    'S8': os.path.join(FIG_PNG, 'Fig3_Decentered_Source_Sensitivity_2026-02-11.png'),
    'S9_city': os.path.join(FIG_PNG, 'Fig2_MI_Sensitivity_Decentered_2026-02-02.png'),
    'S9_nbhd': os.path.join(FIG_PNG, 'Fig3_Decentered_Multiscale_MI_Sensitivity_2026-02-03.png'),
    'S11': os.path.join(FIG_PNG, 'Fig3_Decentered_Multiscale_R6Filter_FullPanel.png'),
    'S12': os.path.join(FIG_PNG, 'Fig3_NudgeSensitivity_FullPanel.png'),
    'S13_pop': os.path.join(FIG_PNG, 'CCDF_City_Population.png'),
    'S13_mass': os.path.join(FIG_PNG, 'CCDF_Total_Built_Mass_tonnes.png'),
    'S13_gi_pop': os.path.join(FIG_PNG, 'RankSize_GI_City_Population.png'),
    'S13_gi_mass': os.path.join(FIG_PNG, 'RankSize_GI_Total_Built_Mass_tonnes.png'),
    'S14': os.path.join(FIG_PNG, 'Fig3_ExtendedData_CityLines_2026-01-30.png'),
    'S15_rank': os.path.join(FIG_PNG, 'RankCorrelation_PopVsMass.png'),
    'S15_perm': os.path.join(FIG_PNG, 'RankPermutation_BetaRobustness.png'),
    'S15_cond': os.path.join(FIG_PNG, 'ConditionalExpectation_LogMass_LogPop.png'),
}

# =============================================================================
# HELPER FUNCTIONS
# =============================================================================

def set_cell_shading(cell, color_hex):
    """Set background shading for a table cell."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_para(doc, text, bold=False, italic=False, size=11, alignment=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Times New Roman'
    if alignment:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(0)
    return p


def add_bold_then_normal(doc, bold_text, normal_text, size=11):
    p = doc.add_paragraph()
    r1 = p.add_run(bold_text)
    r1.bold = True
    r1.font.size = Pt(size)
    r1.font.name = 'Times New Roman'
    r2 = p.add_run(normal_text)
    r2.font.size = Pt(size)
    r2.font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(6)
    return p


def add_figure(doc, fig_key, caption, width=6.0):
    """Insert a figure with caption. Handles missing files gracefully."""
    path = FIGS.get(fig_key)
    if path and os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(width))
        # Caption
        cap = doc.add_paragraph()
        r = cap.add_run(caption)
        r.font.size = Pt(10)
        r.font.name = 'Times New Roman'
        cap.paragraph_format.space_after = Pt(12)
    else:
        p = doc.add_paragraph()
        r = p.add_run(f'[Figure placeholder: {fig_key} — file not found]')
        r.font.size = Pt(10)
        r.italic = True
        cap = doc.add_paragraph()
        r = cap.add_run(caption)
        r.font.size = Pt(10)
        r.font.name = 'Times New Roman'
        cap.paragraph_format.space_after = Pt(12)


def add_table(doc, headers, rows, col_widths=None):
    """Create a formatted table."""
    ncols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, 'D9E2F3')

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = 'Times New Roman'
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # spacing
    return table


# =============================================================================
# DOCUMENT BUILDER
# =============================================================================

def build_si():
    doc = Document()

    # -- Default font --
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    # -- Page setup --
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # =========================================================================
    # TITLE PAGE
    # =========================================================================
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run('Supplementary Information')
    r.bold = True
    r.font.size = Pt(16)
    r.font.name = 'Times New Roman'

    doc.add_paragraph()
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run('Nested economies of scale in city mass')
    r.bold = True
    r.font.size = Pt(14)
    r.font.name = 'Times New Roman'

    authors = doc.add_paragraph()
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = authors.add_run('Kangning Huang, Mingzhen Lu')
    r.font.size = Pt(12)
    r.font.name = 'Times New Roman'

    doc.add_paragraph()
    doc.add_paragraph()

    toc_title = add_para(doc, 'Table of contents', bold=True, size=12)
    toc_items = [
        '1. Materials and Methods',
        '2. Sensitivity Analyses',
        '3. Zipf\'s Law Statistical Tests',
        '4. Neighborhood Variability Analysis',
        '5. Rank Correspondence and Marginalization',
        '6. Interaction Containment and Resolution Dependence',
        '7. Master Confidence Interval Table',
        '8. Figures',
        '9. Tables',
        '10. References',
    ]
    for item in toc_items:
        add_para(doc, f'     {item}', size=11, space_after=3)

    doc.add_page_break()

    # =========================================================================
    # 1. MATERIALS AND METHODS
    # =========================================================================
    add_heading(doc, '1. Materials and Methods', level=1)

    # --- 1.1 Historical trend ---
    add_heading(doc, 'Historical trend of global built structure mass', level=2)
    add_para(doc, (
        'We derived historical global population data from our world in data '
        '(https://ourworldindata.org/grapher/population). We acquired the global built structure mass and global '
        'biomass dataset from Elhacham et al. (2020) from 1900 to 2015 and projection data from 2015 to '
        '2037(1). We used linear interpolation (zoo::na.approx()) to interpolate missing global biomass value '
        'during 1991-1999, 2001-2009, 2011, 2013-2016. To calculate the built structure to human body mass ratio, we '
        'estimated weighted average human body mass as approximately 50 kg. We derived this coarse-grained '
        'estimation by combining the global average age distribution (25% 0-14, 65% 15-64, 10% 65+ in year '
        '2023 based on World Bank data, https://data.worldbank.org/indicator/SP.POP.0014.TO.ZS) with the '
        'age-specific body mass (25kg for 0-14, 60kg for 15-64, 50kg for 65+). We are aware that age-specific '
        'body mass varies across countries, regions, diets, and a range of other factors, however, the qualitative '
        'conclusions presented in our work are not sensitive to the uncertainties associated with this estimation.'
    ))

    # --- 1.2 Vegetation biomass ---
    add_heading(doc, 'Vegetation biomass', level=2)
    add_para(doc, (
        'We extracted city-specific vegetation biomass from a recently published global map of vegetation carbon '
        'that includes both aboveground and belowground biomass carbon (2). The vegetation carbon data is based '
        'on 2010. We then converted biomass carbon into dry-weight vegetation biomass using a convection '
        'coefficient of 0.48g carbon per gram of dry-weight biomass (3). We extracted city-specific vegetation '
        'biomass based on city definitions described in the following section.'
    ))

    # --- 1.3 Definition of cities ---
    add_heading(doc, 'Definition of cities', level=2)
    add_para(doc, (
        'In this study, we define cities using the concept of Urban Cores (UC) provided by the Global Human '
        'Settlement Layer (GHSL). GHSL, developed by the European Commission\'s Joint Research Centre, '
        'delineates urban cores based on continuous areas of built-up land and population density, derived from '
        'satellite imagery and census data(4). An urban core is typically characterized by contiguous built-up '
        'surfaces with high population density, representing the primary agglomeration of urban activity within a '
        'region. Following established conventions, we include only cities with a minimum population size of '
        '50,000 residents in our analysis, ensuring consistency and comparability with other global urban studies '
        '(5, 6). This operational definition allows us to systematically evaluate and compare urban built '
        'environments and associated resource demands across a diverse range of global cities. Our main finding '
        'that economies of scale in urban built stock is prevalent in cities worldwide is robust to the arbitrary '
        'cutoff sizes of cities. In fact, we would argue that the exact cutoff of city definition is not so important '
        'given the self-similarity, or fractal-like feature of urban built stock.'
    ))

    # --- 1.4 Definition of neighborhoods ---
    add_heading(doc, 'Definition of neighborhoods', level=2)
    add_para(doc, (
        'We delineate neighborhoods with the H3 global hexagonal grid (developed by Uber)(7). Hexagons have '
        'two properties that are critical for our analysis. First, among the few regular polygons that tessellate the '
        'plane, hexagons minimize the perimeter-to-area ratio, yielding compact, near-circular units that reduce '
        'edge effects when we aggregate spatial data. Second, hexagons tile space without the directional bias that '
        'characterizes squares (Manhattan effect) or triangles, thereby supporting isotropic distance measures that '
        'better approximate real-world movement patterns. These features mirror the intuition of Christaller\'s '
        'central-place theory(8), which idealizes market areas as hexagons to maximize coverage while '
        'minimizing overlap. Here we use the hexagons at resolution of level-6, where each cell has an average '
        'edge length of ~3.7 km and an area of 37 km\u00b2, a scale that captures the extent of day-to-day human '
        'activity: the cell radius can be covered in ~15 min by bicycle (15 km/h) or ~45 min on foot (5 km/h). The '
        'resulting neighborhoods are therefore large enough to encompass the functional mix of a self-contained '
        'urban district, yet fine-grained enough to retain intra-city variation in the variables we study.'
    ))

    # --- 1.5 Bottom-up approach ---
    add_heading(doc, 'Bottom-up approach of urban built stock', level=2)
    add_para(doc, (
        'To quantify the material stocks embedded in the built environment of cities, we devised a workflow that '
        'is the first global analysis that builds up the total mass of urban buildings and paved surfaces from '
        'the ground up. This method separately estimates the volumes of building and areas of mobility '
        'infrastructure within each city boundary and relies on well-established, type-specific and region-specific '
        'material intensity factors to convert built volume, road surface, and pavements into mass, detailed below.'
    ))

    # 1) Stock of buildings
    add_bold_then_normal(doc,
        '1) Stock of buildings. ',
        'We use raster-based estimates of building volume provided independently by Esch '
        'et al. (2022)(9), Li et al. (2022)(10), and Liu et al. (2024)(11), which offer 3D characterizations of urban '
        'form at various spatial resolutions. To minimize dataset-specific bias, we calculate the average building '
        'volume within each city boundary using all three datasets.'
    )

    # Table S1
    add_bold_then_normal(doc, 'Table S1| ', 'Summaries of input 3D building datasets')
    add_table(doc,
        ['Source', 'Input Imagery (years)', 'Resolution', 'RMSE (volume)'],
        [
            ['Esch et al. (2022)', 'TanDEM-X amplitude (2011, 2013);\nWSF 2019 mask (Sentinel-1/2 pre-2019)', '90 m', '1.54 m\u00b3/m\u00b2'],
            ['Li et al. (2022)', 'Landsat-8 (2015);\nSentinel-1 (2015\u20132016)', '1 km', '0.619 m\u00b3/m\u00b2'],
            ['Liu et al. (2024)', 'Sentinel-1 (2015);\nALOS (2015);\nLandsat-8 (2014-2016)', '500 m', '0.243 m\u00b3/m\u00b2'],
        ]
    )

    add_para(doc, (
        'We then converted building volume to built stock mass using a detailed material intensity framework that '
        'considers both (1) building types and (2) regional construction practices developed by Haberl et al. (2025) '
        'for analyzing global buildings (12). We distinguish building types based on GHS-BUILT-S '
        'classifications, as the distinct function and form of different building types demands different '
        'building standards and consequently different material intensities independent of country and region. '
        'Similarly, material intensity is also a function of country and region due to the dependence of '
        'building on local climate, building practices, developmental level, etc. In Table S2 we provide a '
        'summary look-up table for matching building types to material intensities. For the detailed regional '
        'breakdown (OECD, North America, Japan, China, etc.), please refer to Haberl et al. (2025) Table S5.'
    ))

    # Table S2
    add_bold_then_normal(doc, 'Table S2| ', 'Material intensity as a function of building type and region.')
    add_table(doc,
        ['Building type', 'Building height (m)', 'GHS-BUILT-S (class)', 'Region-specific MI (kg/m\u00b3)'],
        [
            ['Lightweight (LW)', '< 3', 'Residential + non-residential', '151\u2013154'],
            ['Residential single-family house (RS)', '3\u201312', 'Residential', '125\u2013526'],
            ['Residential multi-family house (RM)', '12\u201350', 'Residential', '315\u2013662'],
            ['Non-residential (NR)', '3\u201350', 'Non-residential', '273\u2013654'],
            ['High-rise (HR)', '50\u2013100', 'Residential + non-residential', '312\u2013330'],
        ]
    )

    # 2) Stock of mobility infrastructure
    add_bold_then_normal(doc,
        '2) Stock of mobility infrastructure and other pavement. ',
        'For impervious surface data, we use the '
        'GISA-10m product from Huang et al. (2022)(13), which maps global impervious cover at 10-meter '
        'resolution. To isolate the portion of impervious surface attributable to pavement (e.g., roads, sidewalks), '
        'we subtract the building footprint area\u2014estimated from the previous step by averaging Li et al. '
        '(2022)(10), Liu et al. (2024)(11) and Esch et al. (2022)(9)\u2014from the total impervious surface within '
        'each city.'
    )

    add_para(doc, (
        'After extracting pavement area, we further tease out 5 grades of roads network vs. other pavements '
        '(sidewalks, parking lot) given the big difference in their required structural intensity and consequently '
        'material intensity. For each city, we calculate the total length and total area of the road network per road '
        'grade: with decreasing material intensity, highway, primary, secondary, tertiary, and local. We then '
        'convert road area to road stock mass based on a type-specific and region-specific material intensity '
        'lookup table (summarised in Table S3, full table uploaded as separate file, derived from (14)).'
    ))

    # Table S3 (Road types)
    add_bold_then_normal(doc, 'Table S3| ', 'Road material intensity summary.')
    add_table(doc,
        ['Road Type', '# Lanes (approx.)', 'Lane width (per lane, m)', 'Region- and climate-specific MI (kg/m\u00b2)'],
        [
            ['Highway', '3 \u2013 7', '3.0 \u2013 4.0', '910.6 \u2013 2883.5'],
            ['Primary', '2 \u2013 5', '3.0 \u2013 4.0', '803.9 \u2013 2331.5'],
            ['Secondary', '1 \u2013 5', '2.4 \u2013 4.1', '768.1 \u2013 1817.9'],
            ['Tertiary', '1 \u2013 4', '2.4 \u2013 4.1', '647.7 \u2013 1750.8'],
            ['Local', '1 \u2013 3', '2.4 \u2013 4.1', '538.5 \u2013 1548.2'],
        ]
    )

    add_para(doc, (
        'For other pavements (e.g., sidewalks, parking lot), we assume a material intensity of 570 kg/m\u00b2 '
        '(equivalent to about 20 cm of concrete or asphalt), a value lower than typical road material intensity '
        'adopted by previous meta-analysis in the field of industrial ecology (see Frantz et al. 2023 Table S15 for '
        'parking lots (15)).'
    ))

    add_para(doc, (
        'The total mass of urban built stock is the sum of the estimated building, road mass and other pavement '
        'mass. This bottom-up approach provides a consistent and scalable method to compare material stocks '
        'across global cities using spatially explicit data.'
    ))

    # --- 1.6 Data acquisition for urban population ---
    add_heading(doc, 'Data acquisition for urban population', level=2)
    add_para(doc, (
        'We used the WorldPop dataset to calculate the total population of each city by summing pixels within city '
        'boundaries from WorldPop raster datasets. We chose WorldPop due to its high spatial resolution (100m), '
        'detailed demographic representation, frequent updates, and open accessibility, providing reliable '
        'population data that better captures population distributions compared to other gridded population '
        'datasets(16, 17). WorldPop integrates diverse data sources, including satellite imagery, census, and '
        'ancillary data, making it particularly suitable for accurately mapping populations within urban areas.'
    ))

    # --- 1.7 Scaling theory ---
    add_heading(doc, 'Scaling theory and fixed slope analysis', level=2)
    add_para(doc, (
        'Scaling theory studies the change of system properties as a system changes in its size (18). Its conceptual '
        'foundation can at least be traced back to early insight regarding proportional scaling up (i.e., linear '
        'scaling) built infrastructure in Galileo\'s Two New Sciences (19). Galileo deducted that scaling structures '
        'proportionally in size increases their volume and mass at a faster rate than the increase of strength, '
        'leading to structural failures. Mathematically, this type of relationship is expressed through power-law '
        'equations of the general form:'
    ))
    add_para(doc, 'Y = Y\u2080 N\u1d5d', size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, (
        'In this equation, Y denotes a measurable property or characteristic of the system, N represents system size, '
        '\u03b2 is the scaling exponent capturing how quickly a property changes relative to system size, and Y\u2080 is the '
        'normalization constant, representing the elevation of the scaling relationship on a log-log plot (20, 21), '
        'providing insights into the baseline or developmental level of the system (e.g., a collection of buildings, a '
        'taxon of multiple species, or a group of cities). '
        'In our application, Y is total built material stock (M) and N is population, so M\u2080 \u2261 Y\u2080 represents '
        'the expected material stock at a reference population of N = 1 on the log scale. Because N = 1 lies far '
        'outside the empirical data range, M\u2080 is best interpreted as a baseline parameter capturing the overall '
        'level of material provisioning for a given group of cities: higher M\u2080 indicates that cities in that group '
        'have systematically more built mass at any given population size, reflecting differences in construction '
        'practices, material choices, climate, and economic development.'
    ))
    add_para(doc, (
        'Aside from its widespread application in physics and engineering, scaling theory significantly influenced '
        'biology through the early 20th-century subfield of allometric analysis, enhancing the adoption and '
        'interpretation of Darwin\'s evolutionary theory (22\u201325). Notable examples include scaling relationships '
        'such as organism body mass with metabolic rate and tree height with stem diameter. These foundational '
        'biological concepts were further expanded by West, Brown, and Enquist, who formalized biological '
        'scaling laws, laying the groundwork for applications in other complex systems, including urban contexts '
        'initiated by Bettencourt and colleagues (26).'
    ))
    add_para(doc, (
        'Scaling theory has recently been widely adopted to enhance our understanding of urban systems, '
        'elucidating how a wide range of urban indicators (wealth, innovation, crime, infrastructure, waste, etc.) '
        'scale with city population size (6, 27). Variations in the exponent indicate critical '
        'nonlinearities\u2014economies or diseconomies of scale\u2014that are often obscured by simple per-capita '
        'measures. Superlinear scaling (\u03b2>1) characterizes intensified social interactions and innovation typically '
        'seen in larger cities, whereas sublinear scaling (\u03b2<1) describes constraints related to physical '
        'infrastructure and spatial limits (28).'
    ))
    add_para(doc, (
        'To statistically manage correlations between the scaling exponent \u03b2 and normalization constant Y\u2080, '
        'we apply within-group de-centering (Bettencourt & Lobo, 2016): subtracting group means from both '
        'log-population and log-mass before fitting OLS on the pooled residuals. This approach derives '
        'nation-specific baselines in cross-city analyses and city-specific baselines at neighborhood scales, '
        'isolating the scaling slope from cross-group variation in baseline material intensity.'
    ))

    # --- 1.8 De-centering methodology (NEW - replaces multi-level fixed-effect) ---
    add_heading(doc, 'De-centering methodology', level=2)
    add_para(doc, (
        'We estimate scaling exponents using within-group de-centering, following Bettencourt & Lobo (2016). '
        'For city-level scaling, we subtract each country\'s mean log-population and log-mass from every city, '
        'then fit ordinary least squares (OLS) on the pooled de-centered data:'
    ))
    add_para(doc, (
        'For each country c, define: log_pop_c(i) = log\u2081\u2080(pop_i) \u2212 mean_c[log\u2081\u2080(pop)], and '
        'log_mass_c(i) = log\u2081\u2080(mass_i) \u2212 mean_c[log\u2081\u2080(mass)]. '
        'Then \u03b2 is estimated by OLS: log_mass_c = \u03b2 \u00d7 log_pop_c + \u03b5.'
    ))
    add_para(doc, (
        'For neighborhood-level scaling, we apply the same de-centering logic at the city level. '
        'Let M_ijk denote the built mass of neighborhood k in city j of country i, and N_ijk its population. '
        'We subtract city-specific means:'
    ))
    add_para(doc, (
        'log(M_ijk) \u2212 mean_j[log(M)] = \u03b4 \u00d7 (log(N_ijk) \u2212 mean_j[log(N)]) + \u03b5_ijk'
    ), size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, (
        'where mean_j[\u00b7] denotes the mean over all neighborhoods within city j. This within-city transformation '
        'partials out both country and city baselines, isolating the common within-city scaling slope \u03b4. '
        'Standard errors are clustered at the city level.'
    ))
    add_para(doc, (
        'By the Frisch\u2013Waugh\u2013Lovell theorem, this de-centering estimator is algebraically identical to the '
        'within-group (dummy-variable) fixed-effect estimator. We further validated robustness by fitting the '
        'analogous mixed-effects model with random intercepts u_i and u_j, obtaining virtually identical '
        'estimates of \u03b4. The table below confirms equivalence:'
    ))

    add_table(doc,
        ['', 'De-centering (revised)', 'Mixed-effects (original)'],
        [
            ['City \u03b2', '0.900 [0.890, 0.909]', '0.900 [0.890, 0.910]'],
            ['Neighborhood \u03b4 (R6)', '0.751 [0.747, 0.755]', '0.751 [0.747, 0.755]'],
            ['Neighborhood \u03b4 (R7)', '0.713 [0.712, 0.715]', '0.713 [0.712, 0.715]'],
        ]
    )
    add_para(doc, (
        'Because the outcome and predictor are both in logs, the estimated \u03b4 < 1 directly quantifies sublinearity: '
        'a 1% increase in local population corresponds to only a \u03b4% increase in built mass, uniformly across '
        'all countries and cities.'
    ))

    # --- 1.9 Simulation ---
    add_heading(doc, 'Simulation of the emergence of city-level scaling from neighborhood-level disparity', level=2)
    add_para(doc, (
        'We developed a simulation framework to investigate how variations in the neighborhood-level Zipf\'s law '
        'distribution of population (characterized by the Zipf shape parameter s) influence city-level scaling patterns, '
        'given a fixed neighborhood-level scaling exponent (\u03b4 = 0.75). The total built mass (M) of cities was modeled '
        'as the sum of neighborhood-level built masses (m_i), where m_i = n_i^\u03b4 and n_i is the population of the i-th '
        'neighborhood. With these simulated cities, we then estimated the scaling coefficient (\u03b2) relating M to N '
        'using log-log regression.'
    ))
    add_para(doc, (
        'To explore the emergent relationship between s and \u03b2, we performed a parameter sweep of s values and '
        'analyzed the city-level scaling patterns across a range of scenarios. This approach revealed how the '
        'aggregation of neighborhood-level power-law distributions affects the observed scaling exponent at the '
        'city level. To examine the uncertainties of \u03b2, we introduced two random processes: 1) we drew a random s '
        'parameter from a normal distribution, centered at the given s with the standard deviation (0.59) estimated from '
        'the observed distribution (Fig.3A inset; Fig.4A); 2) we grouped the cities randomly into "countries" by '
        'Dirichlet distribution with shape parameter of 0.1 to generate a power-law like distributions of countries '
        'ranked by numbers of cities. By applying these two random processes, we produce a distribution of theoretical '
        '\u03b2 values.'
    ))

    # --- 1.10 Derivation ---
    add_heading(doc, 'Derivation: city-level scaling law from neighborhood-level disparity', level=2)
    add_para(doc, (
        'We partition each city into K standard hexagonal neighborhoods, labeling them so that n_(1) is the largest '
        'neighborhood population, n_(2) the second largest, and so on and so forth. Likewise, m_(1) is the largest '
        'neighborhood built mass, m_(2) the second largest, etc. Empirical evidence suggests these quantities follow '
        'Zipf-like rank-size laws:'
    ))
    add_para(doc, 'n_(i) ~ i^(\u2212s_n)  and  m_(i) ~ i^(\u2212s_m)', alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, (
        'where the negative exponents s_n and s_m indicate that population and built mass decrease with increasing '
        'rank i. If we posit that the same neighborhood that ranks i-th in population on average also ranks i-th in built '
        'mass, we connect these scalings to obtain:'
    ))
    add_para(doc, 'm_(i) ~ i^(\u2212s_m) = (i^(\u2212s_n))^(s_m/s_n) = (n_(i))^(s_m/s_n)', alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, (
        'Defining a new parameter \u03b4 = s_m / s_n, '
        '\u03b4 is thus the neighborhood-level scaling exponent relating built mass to population: m_(i) ~ n_i^\u03b4. '
        'If s_m < s_n, then \u03b4 < 1, implying sublinear growth of m_(i) with n_(i).'
    ))
    add_para(doc, (
        'Note that in this rank-order formulation, s_n is the neighborhood version of Zipf\'s city rank-order parameter '
        '(31). Empirical evidence suggests that city-level s are frequently found very close to 1 (0.8-1.2). Based on our '
        'analysis, the neighborhood level data (Fig.4A) has a Zipf parameter s_n in the range between 0.7 (25% quantile) '
        'to 1.4 (75% quantile).'
    ))
    add_para(doc, (
        'In our dataset, s_m is consistently smaller than s_n, indicating neighborhood-level mass is more homogenous '
        '(in the limit case, s = 0 represents all neighborhoods are exactly equal in size, the higher the s, the more disparate) '
        'than neighborhood-level population distribution.'
    ))
    add_para(doc, (
        'We then define the total city population N = \u03a3 n_i and total city built mass M = \u03a3 m_i, '
        'and write M ~ N^\u03b2, where \u03b2 is the city-level scaling exponent. If s_n >> 1, the largest population '
        'n_(1) dominates N, and similarly if s_m >> 1 the largest built mass m_(1) dominates M. Then '
        'N \u2248 n_(1), M \u2248 m_(1) ~ (n_(1))^\u03b4, so M ~ N^\u03b4 \u21d2 \u03b4 = \u03b2.'
    ))
    add_para(doc, (
        'Conversely, if s_n is close to zero, so that all neighborhoods are similar in size, then we typically have '
        'N ~ K. If s_m is close to zero as well, we will have M ~ K too, giving: M ~ N \u21d2 \u03b2 = 1.'
    ))
    add_para(doc, (
        'Hence, \u03b2 are bounded between \u03b4 (in a largest-neighborhood-dominated regime) and 1 (in a homogenous '
        'regime).'
    ))

    doc.add_page_break()

    # =========================================================================
    # 2. SENSITIVITY ANALYSES (NEW)
    # =========================================================================
    add_heading(doc, '2. Sensitivity Analyses', level=1)

    # --- 2a. Data source sensitivity ---
    add_heading(doc, '2a. Data source sensitivity', level=2)
    add_para(doc, (
        'Our building volume estimates draw on three independent, globally gridded datasets: '
        'Esch et al. (2022, WSF3D), Li et al. (2022), and Liu et al. (2024, GUS3D). These datasets differ in '
        'spatial resolution (90 m to 1 km), remote sensing platform (radar interferometry, optical/ML ensemble, '
        'and multi-sensor XGBoost), and regional accuracy profiles. To assess how data source choice affects our '
        'scaling exponents, we repeated the full de-centered scaling analysis under five specifications: '
        '(i) each dataset used in isolation; (ii) a random-selection approach drawing one available source per '
        'neighborhood for 100 iterations; and (iii) a reliability-weighted average assigning region-specific weights '
        '(1\u20133 scale) based on each dataset\'s published validation metrics (Table S5).'
    ))
    add_para(doc, (
        'At the city level (Fig. S7, Table S4), \u03b2 ranges from 0.892 (Esch2022) to 0.916 (Liu2024), a total span '
        'of 0.024. The random-selection mean \u03b2 is 0.902 [0.897, 0.907], and the reliability-weighted \u03b2 is '
        '0.901 [0.892, 0.910], both essentially indistinguishable from the baseline (0.899). At the neighborhood '
        'level (Fig. S8, Table S4), \u03b4 ranges from 0.612 (Li2022) to 0.721 (Esch2022). The wider neighborhood-level '
        'range reflects differential spatial coverage: Li2022 provides data for only 90,352 of 141,109 baseline '
        'neighborhoods (64%), so the lower \u03b4 partly reflects a different subsample of cities and neighborhoods '
        'rather than a pure measurement effect. The random-selection approach yields '
        '\u03b4 = 0.714 [0.713, 0.714] with standard deviation < 0.001, confirming that the baseline result is not an '
        'artifact of any single dataset.'
    ))
    add_para(doc, (
        'The reliability-weighted averaging assigns higher weight to the most accurate source in each world region '
        '(Table S5). For example, Liu2024 receives the highest weight in North America (volume R = 0.96) and '
        'China (volume R = 0.84), whereas Li2022 is down-weighted in China (height R\u00b2 = 0.49) but favored in '
        'Oceania (height R\u00b2 = 0.70). Esch2022 provides the only validation data for Central Asia and offers '
        'good low-rise accuracy globally but is penalized for its severe high-rise underestimation (>25 m). '
        'All individual-source and weighted estimates remain sublinear.'
    ))

    # --- 2b. Material intensity sensitivity ---
    add_heading(doc, '2b. Material intensity sensitivity', level=2)
    add_para(doc, (
        'We tested the sensitivity of scaling exponents to material intensity (MI) assumptions using three frameworks '
        'of increasing regional specificity: (1) a global average MI applied uniformly to all cities, (2) a 5-region '
        'differentiation following Haberl et al. (2025), and (3) the 32-region RASMI database from Fishman et al. '
        '(2024), which varies MI by climate zone, development stage, and predominant construction practice. '
        'At the city level, \u03b2 ranges from 0.897 to 0.901 (total range 0.004); at the neighborhood level (Resolution 7), '
        '\u03b4 ranges from 0.706 to 0.713 (total range 0.007, coefficient of variation 0.49%). All estimates are '
        'sublinear and their 95% confidence intervals overlap substantially (Table S6, Fig. S9). This insensitivity '
        'is expected: because MI enters as a multiplicative factor within each region, the de-centering procedure, '
        'which subtracts within-group means, absorbs regional MI differences into the group intercepts, leaving the '
        'slope estimate largely unaffected. These results hold whether MI is specified as a single global value or '
        'differentiated across 32 climate-construction regions.'
    ))

    # --- 2c. Underground infrastructure sensitivity ---
    add_heading(doc, '2c. Underground infrastructure sensitivity', level=2)
    add_para(doc, (
        'Our mass estimates focus on above-ground structures (buildings and roads/pavement). To assess the potential '
        'impact of underground infrastructure, we conducted two analyses: (1) a direct subway sensitivity test and '
        '(2) a magnitude bounding exercise using Frantz et al. (2023).'
    ))
    add_bold_then_normal(doc, 'Subway sensitivity. ',
        'We compiled subway/metro network geometries from CPTOND-2025 (Chinese cities) and OpenStreetMap '
        '(all other countries). We converted these to material mass using two sets of coefficients from Mao et al. '
        '(2021): per-kilometer tunnel mass (~19,500 t/km for concrete, steel, aggregates, and waterproofing) and '
        'per-station mass (~170,000 t/station for the station structure, entrance/exit buildings, ancillary roads, and '
        'ventilation systems). Station mass accounts for roughly 76% of total subway material stocks. Adding subway '
        'mass to baseline estimates shifts city-level \u03b2 from 0.900 to 0.903 (+0.003) and neighborhood-level \u03b4 from '
        '0.713 to 0.714 (+0.001), both well within baseline confidence intervals (Fig. S10). The positive direction '
        'is expected: subway mass concentrates in large, dense cities, slightly increasing the exponent. The negligible '
        'magnitude confirms that subway omission does not affect our findings.'
    )
    add_bold_then_normal(doc, 'Magnitude bounding (Frantz et al. 2023). ',
        'The most spatially explicit material stock assessment to date maps 127 Gt of built structures across the '
        'conterminous US at 10 m resolution. Buildings (including foundations, calibrated via 23 LCA studies) '
        'account for 49% and roads (including base layers) for 39% of total mass. Subways account for only 0.15% '
        'of the total. The components excluded from their assessment\u2014pipelines, electricity networks, '
        'telecommunications\u2014total less than 0.5 t per capita combined, representing <0.1% of the 391 t per capita '
        'in mapped structures (Table S7). Underground utilities are a negligible fraction of total material stocks.'
    )
    add_para(doc, (
        'Even under a theoretical scenario where underground mass comprises 5% of total stocks and scales linearly '
        'with population, the blended exponent would shift by only \u22120.007, remaining firmly sublinear.'
    ))

    # --- 2d. Spatial aggregation sensitivity ---
    add_heading(doc, '2d. Spatial aggregation sensitivity', level=2)
    add_bold_then_normal(doc, 'Multiscale analysis. ',
        'The neighborhood scaling exponent \u03b4 decreases monotonically with finer H3 resolution: from 0.826 at '
        'Resolution 5 (~253 km\u00b2 hexagons) to 0.751 at Resolution 6 (~36 km\u00b2) to 0.713 at Resolution 7 '
        '(~5.2 km\u00b2). This pattern is consistent with the disparity framework presented in Figure 4 of the main text. '
        'At coarser resolution, hexagons average over large areas that include both dense cores and sparse peripheries, '
        'attenuating the within-city population heterogeneity that drives sublinear scaling. As the spatial grain '
        'becomes finer, hexagons capture the full range of neighborhood densities within a city, exposing stronger '
        'economies of scale (lower \u03b4). This resolution dependence does not undermine the sublinear finding. '
        'Instead, it confirms that the scaling relationship operates at every spatial scale examined, with economies '
        'of scale strengthening as the analysis resolves finer-grained urban structure.'
    )
    add_bold_then_normal(doc, 'H3 grid nudge test. ',
        'To test whether the fixed global orientation of the H3 grid introduces systematic bias, we shifted the grid '
        'by 1 km in each cardinal direction and re-estimated \u03b4 at Resolution 7. The maximum deviation across all '
        'four shifts is 0.003 (\u03b4 ranges from 0.713 to 0.716), with all confidence intervals overlapping. The scaling '
        'exponent is insensitive to the arbitrary placement of the hexagonal grid.'
    )

    doc.add_page_break()

    # =========================================================================
    # 3. ZIPF'S LAW STATISTICAL TESTS (NEW)
    # =========================================================================
    add_heading(doc, '3. Zipf\'s Law Statistical Tests', level=1)
    add_para(doc, (
        'We conducted formal distribution testing using two complementary approaches: the Clauset, Shalizi & Newman '
        '(2009) (CSN) framework with Vuong likelihood ratio tests, and the Gabaix & Ibragimov (2011) corrected '
        'rank-size regression.'
    ))
    add_para(doc, (
        'The CSN framework fits a power-law distribution to the upper tail of the data, estimates the lower bound x_min '
        'via KS-distance minimization, and then compares the power-law fit against alternatives (lognormal, exponential, '
        'truncated power law) using the Vuong likelihood ratio test. A positive test statistic R indicates that the '
        'power law is favored over the alternative.'
    ))
    add_para(doc, (
        'The Vuong test strongly favors the power law over the lognormal for both city population (R = 28.9, p < 0.001) '
        'and total built mass (R = 13.9, p < 0.001). The truncated power law (power law with exponential cutoff) '
        'provides the best overall fit for both variables, matching the empirical consensus that city-size distributions '
        'follow power-law behavior in the upper tail with a soft cutoff at the largest sizes.'
    ))
    add_para(doc, (
        'The Gabaix-Ibragimov rank-size regression yields Zipf exponents of \u03b6_pop = 0.946 [0.903, 0.988] and '
        '\u03b6_mass = 0.877 [0.838, 0.917]. Exact Zipf\'s law (\u03b6 = 1) is statistically rejected for both, but both '
        'distributions are "near-Zipf" with high R\u00b2 (> 0.93). The lower mass exponent follows directly from '
        'sublinear scaling (\u03b2 < 1): if population is near-Zipf with exponent \u03b6_pop, then mass with '
        'M ~ N^\u03b2 implies \u03b6_mass \u2248 \u03b2 \u00d7 \u03b6_pop. Our theoretical framework requires heavy-tailed '
        'distributions with a characteristic exponent, not exact Zipf, and these tests confirm that requirement.'
    ))

    doc.add_page_break()

    # =========================================================================
    # 4. NEIGHBORHOOD VARIABILITY ANALYSIS (NEW)
    # =========================================================================
    add_heading(doc, '4. Neighborhood Variability Analysis', level=1)
    add_para(doc, (
        'Does the variability in city-specific neighborhood scaling exponents undermine the universality claim? '
        'We argue it does not, for three reasons.'
    ))
    add_bold_then_normal(doc, 'First, variation is a theoretical prediction. ',
        'Scaling theory explicitly predicts that individual cities deviate from the mean relationship, with residuals '
        'encoding local urban characteristics such as urban form, infrastructure age, and development history '
        '(Bettencourt et al., 2010; Gomez-Lievano et al., 2016; Lobo et al., 2025). The universality claim concerns '
        'the population-level average, not the trajectory of every individual city.'
    )
    add_bold_then_normal(doc, 'Second, the data overwhelmingly support sublinear scaling. ',
        'Across 3,312 cities in 72 countries, only 1 has a negative neighborhood scaling exponent (99.97% positive), '
        'with mean slope 0.754 and median 0.764 (Table S10). The interquartile range is [0.66, 0.86], indicating '
        'that the vast majority of cities exhibit sublinear scaling at the neighborhood level. The pooled R\u00b2 is 0.86.'
    )
    add_bold_then_normal(doc, 'Third, our methodology eliminates MAUP. ',
        'The H3 hexagonal grid system assigns each neighborhood a globally standardized area and shape, '
        'removing the arbitrariness of administratively defined neighborhoods. This means the observed variation '
        'reflects genuine differences in urban structure\u2014not boundary artifacts.'
    )

    doc.add_page_break()

    # =========================================================================
    # 5. RANK CORRESPONDENCE AND MARGINALIZATION (NEW)
    # =========================================================================
    add_heading(doc, '5. Rank Correspondence and Marginalization', level=1)
    add_para(doc, (
        'Our regression does not assume that the city ranked k-th in population is also ranked k-th in material stock. '
        'The scaling exponent is estimated from paired observations (log N_i, log M_i) for each city i via OLS on '
        'de-centered data. The exponent therefore reflects the conditional expectation E[log M | log N], consistent '
        'with the marginalization framework of urban scaling analysis (Bettencourt et al., 2012; Ribeiro et al., 2021).'
    ))
    add_para(doc, (
        'To provide empirical support, we computed rank-order statistics and performed permutation tests '
        '(Fig. S15, Table S11). The Spearman rank correlation between population and material stock is '
        '\u03c1 = 0.771 (Kendall \u03c4 = 0.581), confirming substantial but imperfect rank correspondence. '
        'A within-decile permutation test\u2014in which mass values are shuffled among cities within the same '
        'population decile\u2014yields \u03b2 = 0.835 (vs. observed \u03b2 = 0.900), and full random shuffling collapses '
        '\u03b2 to zero. The nonparametric conditional expectation (30 quantile bins) matches the OLS estimate '
        'within 0.001 (R\u00b2 = 0.999). These diagnostics confirm that the scaling relationship is a genuine '
        'structural association between population and material stock, not a statistical artifact of rank ordering.'
    ))

    doc.add_page_break()

    # =========================================================================
    # 6. INTERACTION CONTAINMENT (NEW)
    # =========================================================================
    add_heading(doc, '6. Interaction Containment and the Resolution Dependence of Scaling', level=1)
    add_para(doc, (
        'Standard urban scaling theory assumes that all relevant social and economic interactions occur within the '
        'city boundary\u2014that the city captures 100% of the interaction volume driving infrastructure demand '
        '(Bettencourt, 2013). At sub-city scales, this assumption breaks down: a neighborhood hexagon captures '
        'only a fraction of resident interactions, with the remainder spilling to adjacent areas.'
    ))
    add_para(doc, (
        'Two physical mechanisms produce the sublinear scaling of material stocks. First, vertical densification '
        'reduces per-capita material requirements: larger and denser cities tend to have taller buildings, and '
        'multi-story buildings share foundations, roofs, and structural cores across many dwelling units, lowering '
        'the material cost per occupant. Second, infrastructure sharing generates savings in mobility networks: '
        'in denser settlements, a given length of road serves more residents, so network length grows more slowly '
        'than population. These mechanisms operate jointly\u2014taller buildings both house more people per unit '
        'land and concentrate demand for shared infrastructure.'
    ))
    add_para(doc, (
        'We quantify the spatial reach of infrastructure sharing using the truncated power-law displacement '
        'distribution reported by Gonz\u00e1lez, Hidalgo & Barab\u00e1si (2008), which characterizes individual human '
        'mobility as: P(\u0394r) \u221d (\u0394r + \u0394r\u2080)^(\u2212(1+\u03b2_m)) \u00b7 exp(\u2212\u0394r/\u03ba), '
        'with \u03b2_m = 1.75, \u0394r\u2080 = 1.5 km, and \u03ba = 400 km (D1 dataset parameters). Evaluating the '
        'cumulative distribution F(R) = Pr(\u0394r \u2264 R) at the edge lengths of our H3 hexagons via numerical '
        'quadrature yields:'
    ))

    add_table(doc,
        ['H3 resolution', 'Edge length R (km)', 'Hex area (km\u00b2)', 'F(R): share of steps \u2264 R'],
        [
            ['7', '1.4', '5.2', '68.7%'],
            ['6', '3.7', '36.1', '88.9%'],
            ['5', '9.8', '252.9', '97.2%'],
        ]
    )

    add_para(doc, (
        'At our finest analysis scale (Resolution 7, ~5 km\u00b2), approximately 69% of displacement steps fall '
        'within a single hexagon radius, meaning that most routine activities (commuting, shopping, social visits) '
        'occur locally, and approximately 31% of movements cross hexagon boundaries.'
    ))
    add_para(doc, (
        'This interaction containment framework explains two key empirical patterns:'
    ))
    add_para(doc, (
        '1. \u03b4 < \u03b2: Because neighborhoods capture less than 100% of interactions, the economies of scale from '
        'infrastructure sharing are partially diluted\u2014some shared infrastructure serves residents who interact '
        'across neighborhood boundaries. City boundaries, by contrast, contain virtually all routine interactions, '
        'yielding a larger (less sublinear) exponent.'
    ))
    add_para(doc, (
        '2. \u03b4 increases monotonically with coarser spatial resolution (from 0.713 at Resolution 7 to 0.751 at '
        'Resolution 6 to 0.826 at Resolution 5): as hexagons grow, F(R) rises from 69% to 89% to 97%, a larger '
        'share of interactions is contained, and the apparent economies of scale weaken accordingly, approaching '
        'the city-level exponent as the spatial unit approaches city scale.'
    ))
    add_para(doc, (
        'F(R) represents step-length containment (the share of consecutive displacements shorter than R), not exact '
        'geometric containment within a hexagonal cell. It provides a theoretically grounded, first-order proxy '
        'connecting our multiscale empirical results to the established literature on human mobility.'
    ))

    doc.add_page_break()

    # =========================================================================
    # 7. MASTER CI TABLE (NEW)
    # =========================================================================
    add_heading(doc, '7. Master Confidence Interval Table', level=1)
    add_para(doc, (
        'Table S12 compiles 95% confidence intervals for all primary and sensitivity analyses. Across every '
        'specification\u2014varying the building volume data source, material intensity framework, inclusion of subway '
        'infrastructure, spatial resolution, and H3 grid placement\u2014the city-level exponent \u03b2 remains sublinear '
        '(range 0.892\u20130.916) and the neighborhood-level exponent \u03b4 ranges from 0.612 to 0.721. The confidence '
        'intervals for \u03b2 and \u03b4 never overlap in any analysis. The minimum gap between the city CI lower bound '
        'and the neighborhood CI upper bound is 0.16, corresponding to z = 39.1 (p < 0.001) for the primary '
        'comparison. The distinction between city-level and neighborhood-level scaling is therefore an empirical '
        'regularity, not an artifact of any particular methodological choice.'
    ))

    doc.add_page_break()

    # =========================================================================
    # 8. FIGURES
    # =========================================================================
    add_heading(doc, '8. Figures', level=1)

    # --- Existing Figures S1-S6 (from original SI pages) ---
    add_figure(doc, 'S1',
        'Fig. S1| Illustration of city definition, example cities across various continents demonstrating '
        'different scales of cities.',
        width=6.5)

    doc.add_page_break()
    add_figure(doc, 'S2',
        'Fig. S2| Sublinear urban scaling of built mass and variations across all countries. We have 74 '
        'countries meeting our criteria of inclusion (n_cities>5).',
        width=6.5)

    doc.add_page_break()
    add_figure(doc, 'S3',
        'Fig. S3. Country-level scaling analysis stats. The top row (A,B,C) includes all 74 countries with more '
        'than 5 cities, while the second row (D,E,F) excludes 3 countries with R\u00b2 lower than 70%. The overall '
        'country-level stats are not sensitive to the inclusion of these 3 countries with relatively low goodness of fit.',
        width=6.5)

    doc.add_page_break()
    add_figure(doc, 'S4',
        'Fig. S4. Correlation matrix of building mass estimates across three global datasets. The heatmap '
        'shows Pearson correlation coefficients between city-level building mass estimates from Esch2022, '
        'Li2022, and Liu2024 datasets. Analysis includes cities with complete data across all three datasets (n=3,721).',
        width=4.5)

    doc.add_page_break()
    add_figure(doc, 'S5',
        'Fig. S5: Distribution of coefficient of variation (CV) of urban building mass estimates across cities '
        'from multiple data sources. This histogram illustrates the distribution of the CV of urban building mass '
        'estimates across cities, calculated from three global building volume datasets: Esch et al. (2022), '
        'Li et al. (2022), and Liu et al. (2024). The red dashed line indicates the median CV, while the orange '
        'dashed line represents the mean CV.',
        width=5.5)

    doc.add_page_break()
    add_figure(doc, 'S6',
        'Fig. S6: Coefficient of variation in urban building mass estimates for the 30 most populous cities '
        'with multi-source data. Cities are ranked by CV magnitude, showing the relative variability in building '
        'mass estimates between the three global datasets.',
        width=5.0)

    doc.add_page_break()

    # --- NEW Figures S7-S15 ---
    add_figure(doc, 'S7',
        'Fig. S7| City-level data source sensitivity. De-centered scaling analysis repeated with each of three '
        'building volume datasets individually, with random per-city source selection (100 iterations), and with '
        'reliability-weighted averaging. City-level \u03b2 ranges from 0.892 to 0.916 (total span 0.024), all sublinear.',
        width=6.0)

    doc.add_page_break()
    add_figure(doc, 'S8',
        'Fig. S8| Neighborhood-level data source sensitivity. De-centered scaling analysis repeated with each '
        'building volume dataset individually at the neighborhood level. \u03b4 ranges from 0.612 (Li2022) to 0.721 '
        '(Esch2022). All estimates are sublinear regardless of data source choice.',
        width=6.0)

    doc.add_page_break()
    add_figure(doc, 'S9_city',
        'Fig. S9a| City-level material intensity sensitivity. Comparison of scaling exponents under three MI '
        'frameworks: global average, Haberl 5-region, and Fishman 32-region RASMI. \u03b2 ranges from 0.897 to 0.901 '
        '(total range 0.004).',
        width=6.0)
    add_figure(doc, 'S9_nbhd',
        'Fig. S9b| Neighborhood-level material intensity sensitivity. \u03b4 ranges from 0.706 to 0.713 '
        '(total range 0.007, CV = 0.49%). The insensitivity arises because MI acts as a multiplicative constant '
        'absorbed by de-centering.',
        width=6.0)

    doc.add_page_break()
    # Fig. S10 placeholder (subway figures not yet generated)
    p = doc.add_paragraph()
    r = p.add_run('[Fig. S10 — Subway infrastructure impact on scaling: Figure to be generated from '
                   'Fig2_UniversalScaling_Decentered_WithSubwayMass.R and '
                   'Fig3_NeighborhoodScaling_Decentered_R7_WithSubwayMass.R]')
    r.font.size = Pt(10)
    r.italic = True
    r.font.color.rgb = RGBColor(150, 0, 0)
    cap = doc.add_paragraph()
    r = cap.add_run(
        'Fig. S10| Subway infrastructure impact on scaling. Adding subway/metro infrastructure mass '
        '(tunnels and stations, from Mao et al. 2021) shifts city-level \u03b2 by +0.003 and neighborhood-level '
        '\u03b4 by +0.001, both well within baseline confidence intervals.')
    r.font.size = Pt(10)

    doc.add_page_break()
    add_figure(doc, 'S11',
        'Fig. S11| Multiscale sensitivity analysis (H3 Resolutions 5, 6, and 7). The neighborhood scaling '
        'exponent \u03b4 decreases monotonically with finer resolution: 0.826 (R5) \u2192 0.751 (R6) \u2192 0.713 (R7), '
        'but remains sublinear at every scale. Consistent filtering criteria applied across all resolutions.',
        width=6.0)

    doc.add_page_break()
    add_figure(doc, 'S12',
        'Fig. S12| H3 nudge sensitivity. The H3 grid was shifted by 1 km in each of four cardinal directions '
        '(East, North, South, West) and the scaling analysis re-run at Resolution 7. Maximum \u03b4 deviation '
        'is 0.003 (range: 0.713\u20130.716), with fully overlapping confidence intervals.',
        width=6.0)

    doc.add_page_break()
    add_figure(doc, 'S13_pop',
        'Fig. S13a| Complementary cumulative distribution function (CCDF) of city population. '
        'Empirical data (circles) fitted with power law (red), lognormal (blue), and truncated power law (green). '
        'The truncated power law provides the best fit (Vuong test R = \u221272.3, p < 0.001 vs. simple power law).',
        width=5.0)
    add_figure(doc, 'S13_mass',
        'Fig. S13b| CCDF of total built mass (tonnes). The truncated power law again provides the best fit. '
        'Vuong test: power law vs. lognormal R = 13.9, p < 0.001.',
        width=5.0)

    doc.add_page_break()
    add_figure(doc, 'S13_gi_pop',
        'Fig. S13c| Gabaix-Ibragimov rank-size regression for city population. '
        '\u03b6_pop = 0.946 [0.903, 0.988], R\u00b2 = 0.991. Near-Zipf but exact Zipf (\u03b6 = 1) rejected (p = 0.012).',
        width=5.0)
    add_figure(doc, 'S13_gi_mass',
        'Fig. S13d| Gabaix-Ibragimov rank-size regression for total built mass. '
        '\u03b6_mass = 0.877 [0.838, 0.917], R\u00b2 = 0.935. The lower mass exponent is consistent with sublinear scaling.',
        width=5.0)

    doc.add_page_break()
    add_figure(doc, 'S14',
        'Fig. S14| All 3,312 city-specific neighborhood regression lines overlaid. Each thin line represents '
        'the within-city OLS regression of de-centered log-mass on de-centered log-population. Only 1 of 3,312 '
        'cities shows a negative slope (99.97% positive). Mean slope = 0.754, median = 0.764.',
        width=6.0)

    doc.add_page_break()
    add_figure(doc, 'S15_rank',
        'Fig. S15a| Rank-rank scatter plot of city population rank vs. material stock rank. '
        'Spearman \u03c1 = 0.771, Kendall \u03c4 = 0.581, confirming substantial but imperfect rank correspondence.',
        width=5.0)
    add_figure(doc, 'S15_perm',
        'Fig. S15b| Permutation test for \u03b2 robustness. Within-decile shuffling yields \u03b2 = 0.835 '
        '(vs. observed 0.900). Full random shuffling collapses \u03b2 to zero, confirming the structural association.',
        width=5.0)
    add_figure(doc, 'S15_cond',
        'Fig. S15c| Nonparametric conditional expectation E[log M | log N] computed from 30 quantile bins. '
        'The conditional expectation slope matches the OLS estimate within 0.001 (R\u00b2 = 0.999).',
        width=5.0)

    doc.add_page_break()

    # =========================================================================
    # 9. TABLES
    # =========================================================================
    add_heading(doc, '9. Tables', level=1)

    # --- Table S4: Data source sensitivity results ---
    add_bold_then_normal(doc, 'Table S4| ', 'Data source sensitivity results with 95% confidence intervals.')
    add_table(doc,
        ['Level', 'Specification', 'Exponent', '95% CI', 'R\u00b2', 'N'],
        [
            ['City (\u03b2)', 'Esch2022 only', '0.892', '[0.882, 0.902]', '0.900', '3,588'],
            ['City (\u03b2)', 'Li2022 only', '0.899', '[0.890, 0.909]', '0.904', '3,588'],
            ['City (\u03b2)', 'Liu2024 only', '0.916', '[0.906, 0.926]', '0.902', '3,588'],
            ['City (\u03b2)', 'Random selection (mean)', '0.902', '[0.897, 0.907]', '\u2014', '3,588'],
            ['City (\u03b2)', 'Reliability-weighted', '0.901', '[0.892, 0.910]', '0.907', '3,588'],
            ['City (\u03b2)', 'Baseline (3-source avg)', '0.900', '[0.890, 0.909]', '0.909', '3,588'],
            ['Nbhd (\u03b4)', 'Esch2022 only', '0.721', '[0.719, 0.723]', '0.867', '125,431'],
            ['Nbhd (\u03b4)', 'Li2022 only', '0.612', '[0.609, 0.615]', '0.832', '90,352'],
            ['Nbhd (\u03b4)', 'Liu2024 only', '0.700', '[0.698, 0.702]', '0.854', '133,880'],
            ['Nbhd (\u03b4)', 'Random selection (mean)', '0.714', '[0.713, 0.714]', '\u2014', '141,109'],
            ['Nbhd (\u03b4)', 'Reliability-weighted', '0.620', '[0.617, 0.623]', '0.841', '86,503'],
            ['Nbhd (\u03b4)', 'Baseline (3-source avg)', '0.713', '[0.712, 0.715]', '0.862', '141,109'],
        ]
    )

    # --- Table S5: Regional accuracy comparison ---
    add_bold_then_normal(doc, 'Table S5| ', 'Regional accuracy comparison across three building volume datasets. '
        'Ratings: HIGH = strong validation + good metrics; MODERATE = some validation; LOW = limited validation; '
        'UNKNOWN = no validation data.')
    add_table(doc,
        ['World Region', 'Esch2022', 'Li2022', 'Liu2024', 'Most Reliable'],
        [
            ['North America (USA)', 'MOD-HIGH', 'HIGH', 'HIGH', 'Liu2024'],
            ['Canada', 'No validation', 'HIGH', 'HIGH', 'Li / Liu'],
            ['Europe (W. & C.)', 'MOD-HIGH', 'HIGH', 'HIGH', 'Liu2024'],
            ['East Asia (China)', 'MODERATE', 'LOW-MOD', 'MOD-HIGH', 'Liu2024'],
            ['East Asia (Japan/Korea)', 'LOW-MOD', 'MODERATE', 'MODERATE', 'Mixed'],
            ['South Asia', 'MODERATE', 'LOW-MOD', 'LOW', 'Esch2022'],
            ['Southeast Asia', 'MODERATE', 'MODERATE', 'LOW', 'Li2022'],
            ['Sub-Saharan Africa', 'MODERATE', 'MODERATE', 'LOW', 'Esch2022'],
            ['Latin America', 'LOW-MOD', 'MODERATE', 'LOW', 'Li2022'],
            ['Oceania', 'LOW-MOD', 'HIGH', 'MODERATE', 'Li2022'],
            ['MENA', 'MODERATE', 'MODERATE', 'UNKNOWN', 'Esch / Li'],
            ['Central Asia', 'LOW', 'LOW', 'UNKNOWN', 'Esch2022'],
            ['Russia / N. Asia', 'No validation', 'LOW-MOD', 'LOW-MOD', 'Liu2024'],
        ]
    )

    # --- Table S6: MI sensitivity summary ---
    add_bold_then_normal(doc, 'Table S6| ', 'Material intensity sensitivity summary.')
    add_table(doc,
        ['Level', 'MI Approach', 'Exponent', '95% CI', 'R\u00b2', 'N'],
        [
            ['City (\u03b2)', 'Global Average', '0.897', '[0.887, 0.906]', '0.908', '3,588'],
            ['City (\u03b2)', 'Haberl 5-region (baseline)', '0.899', '[0.890, 0.909]', '0.909', '3,588'],
            ['City (\u03b2)', 'Fishman RASMI 32-region', '0.901', '[0.891, 0.911]', '0.909', '3,588'],
            ['Nbhd (\u03b4, R7)', 'Global Average', '0.706', '[0.705, 0.708]', '0.862', '141,109'],
            ['Nbhd (\u03b4, R7)', 'Haberl 5-region (baseline)', '0.713', '[0.712, 0.715]', '0.862', '141,109'],
            ['Nbhd (\u03b4, R7)', 'Fishman RASMI 32-region', '0.709', '[0.708, 0.711]', '0.860', '141,109'],
        ]
    )

    # --- Table S7: Underground infrastructure bounding ---
    add_bold_then_normal(doc, 'Table S7| ', 'Underground infrastructure: component breakdown and impact on scaling.')
    add_table(doc,
        ['Component', 'Source', 'Magnitude', 'Share of total', 'Impact on scaling'],
        [
            ['Buildings (incl. foundations)', 'Frantz et al. 2023', '62.0 Gt (US)', '49%', 'Included in baseline'],
            ['Roads (incl. base layers)', 'Frantz et al. 2023', '~50 Gt (US)', '~39%', 'Included in baseline'],
            ['Subways', 'Frantz et al.; this study', '~0.19 Gt (US)', '0.15%', '\u03b2 +0.003, \u03b4 +0.001'],
            ['Electricity networks', 'Frantz et al. 2023', '~0.15 Gt (US)', '0.12%', 'Spatially uniform'],
            ['Pipelines (water/sewer/gas)', 'Frantz et al. 2023', '<0.5 t/cap', '<0.1%', 'Spatially uniform'],
            ['All excluded components', 'Frantz et al. 2023', '<0.5% of total', '<0.5%', 'Cannot bias slope'],
        ]
    )

    # --- Table S8: Multiscale and nudge summary ---
    add_bold_then_normal(doc, 'Table S8| ', 'Multiscale and H3 nudge sensitivity summary.')
    add_table(doc,
        ['Analysis', 'Specification', '\u03b4', '95% CI', 'R\u00b2', 'N nbhds', 'N cities'],
        [
            ['Multiscale', 'Resolution 5', '0.826', '[0.814, 0.839]', '0.953', '837', '45'],
            ['Multiscale', 'Resolution 6', '0.751', '[0.747, 0.755]', '0.893', '17,497', '735'],
            ['Multiscale', 'Resolution 7 (primary)', '0.713', '[0.712, 0.715]', '0.862', '141,109', '3,312'],
            ['Nudge (R7)', 'Original', '0.713', '[0.712, 0.715]', '0.862', '141,109', '3,312'],
            ['Nudge (R7)', 'East +1 km', '0.715', '[0.712, 0.717]', '0.746', '139,977', '3,312'],
            ['Nudge (R7)', 'North +1 km', '0.713', '[0.711, 0.716]', '0.747', '140,288', '3,312'],
            ['Nudge (R7)', 'South +1 km', '0.716', '[0.714, 0.718]', '0.742', '140,218', '3,312'],
            ['Nudge (R7)', 'West +1 km', '0.716', '[0.714, 0.718]', '0.755', '140,372', '3,312'],
        ]
    )

    # --- Table S9: Zipf's law ---
    add_bold_then_normal(doc, 'Table S9| ', 'Zipf\'s law statistical tests.')
    add_para(doc, 'Panel A: Clauset-Shalizi-Newman framework with Vuong likelihood ratio tests.', italic=True, size=10)
    add_table(doc,
        ['Variable', '\u03b1 (SE)', 'x_min', 'N_tail', 'KS', 'PL vs LN: R (p)', 'PL vs Exp: R (p)', 'PL vs TPL: R (p)'],
        [
            ['City Population', '1.895 (0.015)', '50,044', '3,756', '0.014', '28.9 (<0.001)', '19.9 (<0.001)', '\u221272.3 (<0.001)'],
            ['Total Built Mass', '2.134 (0.038)', '70.8 Mt', '885', '0.024', '13.9 (<0.001)', '16.3 (<0.001)', '\u221228.2 (<0.001)'],
        ]
    )
    add_para(doc, 'Panel B: Gabaix-Ibragimov corrected rank-size regression.', italic=True, size=10)
    add_table(doc,
        ['Variable', '\u03b6', '95% CI', 'R\u00b2', 'H\u2080: \u03b6 = 1'],
        [
            ['City Population', '0.946', '[0.903, 0.988]', '0.991', 'Reject (z = \u22122.50, p = 0.012)'],
            ['Total Built Mass', '0.877', '[0.838, 0.917]', '0.935', 'Reject (z = \u22126.07, p < 0.001)'],
        ]
    )

    # --- Table S10: City-specific slope distribution ---
    add_bold_then_normal(doc, 'Table S10| ', 'Distribution of city-specific neighborhood scaling slopes (3,312 cities).')
    add_table(doc,
        ['Statistic', 'Value'],
        [
            ['N (cities)', '3,312'],
            ['N (countries)', '72'],
            ['Mean slope', '0.754'],
            ['Median slope', '0.764'],
            ['Standard deviation', '~0.14'],
            ['Interquartile range', '[0.66, 0.86]'],
            ['Minimum', '< 0 (1 city)'],
            ['Maximum', '~1.3'],
            ['% positive slopes', '99.97%'],
            ['Pooled R\u00b2', '0.86'],
        ]
    )

    # --- Table S11: Rank correlation and permutation tests ---
    add_bold_then_normal(doc, 'Table S11| ', 'Rank correlation and permutation test results.')
    add_table(doc,
        ['Statistic', 'Value'],
        [
            ['Spearman \u03c1 (population vs. mass rank)', '0.771'],
            ['Kendall \u03c4', '0.581'],
            ['Mean absolute rank displacement', '549 positions'],
            ['Observed \u03b2 (OLS on de-centered data)', '0.900'],
            ['Within-decile shuffle \u03b2 (mean \u00b1 SD, 1000 iter)', '0.835 \u00b1 0.003'],
            ['Full random shuffle \u03b2 (mean \u00b1 SD, 1000 iter)', '\u22120.001 \u00b1 0.016'],
            ['Conditional expectation slope (30 bins)', '0.900'],
            ['Conditional expectation R\u00b2 vs. OLS fit', '0.999'],
        ]
    )

    # --- Table S12: Master CI table ---
    add_bold_then_normal(doc, 'Table S12| ', 'Master confidence interval table: all primary and sensitivity analyses.')
    add_table(doc,
        ['Panel', 'Analysis', 'Exponent', '95% CI', 'CIs overlap?'],
        [
            ['A', 'Primary: City', '\u03b2 = 0.900', '[0.890, 0.909]', 'No (gap = 0.175)'],
            ['A', 'Primary: Neighborhood (R7)', '\u03b4 = 0.713', '[0.712, 0.715]', ''],
            ['B', 'Data source, city (range)', '\u03b2: 0.892\u20130.916', 'All sublinear', 'N/A'],
            ['C', 'Data source, nbhd (range)', '\u03b4: 0.612\u20130.721', 'All sublinear', 'N/A'],
            ['D', 'MI, city (range)', '\u03b2: 0.897\u20130.901', '[0.887, 0.911]', 'N/A'],
            ['E', 'MI, nbhd (range)', '\u03b4: 0.706\u20130.713', '[0.705, 0.715]', 'N/A'],
            ['F', 'Subway: City', '\u03b2 = 0.903', '+0.003 from baseline', 'No'],
            ['F', 'Subway: Neighborhood', '\u03b4 = 0.714', '+0.001 from baseline', ''],
            ['G', 'Multiscale: R5', '\u03b4 = 0.826', '[0.814, 0.839]', 'N/A'],
            ['G', 'Multiscale: R6', '\u03b4 = 0.751', '[0.747, 0.755]', ''],
            ['G', 'Multiscale: R7', '\u03b4 = 0.713', '[0.712, 0.715]', ''],
            ['H', 'Nudge R7 (range)', '\u03b4: 0.713\u20130.716', 'All overlapping', 'N/A'],
        ]
    )

    # --- Original SI Tables 1-3 (data files) ---
    doc.add_paragraph()
    add_bold_then_normal(doc, 'SI Table 1| ',
        'Global summary of various built stock vs. human, vegetation table (source data for figure 1c). '
        'Submitted as a separate SI file (Fig1_BoxplotData_20250705.csv), also available on figshare depository.')
    add_bold_then_normal(doc, 'SI Table 2| ',
        'By country scaling analysis slope and intercept. '
        'Submitted as a separate SI file (country_summary_table.csv), also available on figshare depository.')
    add_bold_then_normal(doc, 'SI Table 3| ',
        'By city scaling analysis slope and intercept (source data for figure 3). '
        'Submitted as a separate SI file (city_inset_slope_2025-07-05.csv), also available on figshare depository.')

    doc.add_page_break()

    # =========================================================================
    # 10. REFERENCES
    # =========================================================================
    add_heading(doc, '10. References', level=1)

    refs = [
        'E. Elhacham, L. Ben-Uri, J. Grozovski, Y. M. Bar-On, R. Milo, Global human-made mass exceeds all living biomass. Nature 588, 442\u2013444 (2020).',
        'S. A. Spawn, C. C. Sullivan, T. J. Lark, H. K. Gibbs, Harmonized global maps of above and belowground biomass carbon density in the year 2010. Sci. Data 7, 112 (2020).',
        'M. Lu, L. O. Hedin, Global plant\u2013symbiont organization and emergence of biogeochemical cycles resolved by evolution-based trait modelling. Nature Ecology & Evolution 3, 239\u2013250 (2019).',
        'M. Melchiorri, M. Pesaresi, A. J. Florczyk, C. Corbane, T. Kemper, Principles and applications of the Global Human Settlement Layer as baseline for the Land Use Efficiency indicator\u2014SDG 11.3.1. ISPRS Int. J. Geoinf. 8, 96 (2019).',
        'A. Cattaneo, S. Girgin, R. de By, T. McMenomy, A. Nelson, S. Vaz, Worldwide delineation of multi-tier city\u2013regions. Nat Cities 1, 469\u2013479 (2024).',
        'M. Lu, C. Zhou, C. Wang, R. B. Jackson, C. P. Kempes, Worldwide scaling of waste generation in urban systems. Nat Cities 1, 126\u2013135 (2024).',
        'J. Bousquin, Discrete Global Grid Systems as scalable geospatial frameworks for characterizing coastal environments. Environ. Model. Softw. 146, 1\u201314 (2021).',
        'B. Goodall, The Penguin Dictionary of Human Geography (Penguin Books, Harlow, England, 1987).',
        'T. Esch, E. Brzoska, S. Dech, B. Leutner, D. Palacios-Lopez, A. Metz-Marconcini, M. Marconcini, A. Roth, J. Zeidler, World Settlement Footprint 3D - A first three-dimensional survey of the global building stock. Remote Sens. Environ. 270, 112877 (2022).',
        'M. Li, Y. Wang, J. F. Rosier, P. H. Verburg, J. van Vliet, Global maps of 3D built-up patterns for urban morphological analysis. Int. J. Appl. Earth Obs. Geoinf. 114, 103048 (2022).',
        'X. Liu, X. Wu, X. Li, X. Xu, W. Liao, L. Jiao, Z. Zeng, G. Chen, X. Li, Global mapping of three-dimensional (3D) urban structures reveals escalating utilization in the vertical dimension and pronounced building space inequality. Engineering (Beijing), doi: 10.1016/j.eng.2024.01.025 (2024).',
        'H. Haberl, A. Baumgart, J. Zeidler, F. Schug, D. Frantz, D. Palacios-Lopez, T. Fishman, Y. Peled, B. Cai, D. Vir\u00e1g, P. Hostert, D. Wiedenhofer, T. Esch, Weighing the global built environment: High-resolution mapping and quantification of material stocks in buildings. J. Ind. Ecol. 29, 159\u2013172 (2025).',
        'X. Huang, J. Yang, W. Wang, Z. Liu, Mapping 10 m global impervious surface area (GISA-10m) using multi-source geospatial data. Earth Syst. Sci. Data 14, 3649\u20133672 (2022).',
        'L. S. A. Rousseau, B. Kloostra, H. AzariJafari, S. Saxe, J. Gregory, E. G. Hertwich, Material stock and embodied greenhouse gas emissions of global and urban road pavement. Environ. Sci. Technol. 56, 18050\u201318059 (2022).',
        'D. Frantz, F. Schug, D. Wiedenhofer, A. Baumgart, D. Vir\u00e1g, S. Cooper, C. G\u00f3mez-Medina, F. Lehmann, T. Udelhoven, S. van der Linden, P. Hostert, H. Haberl, Unveiling patterns in human dominated landscapes through mapping the mass of US built structures. Nat. Commun. 14, 8014 (2023).',
        'D. R. Thomson, D. R. Leasure, T. Bird, N. Tzavidis, A. J. Tatem, How accurate are WorldPop-Global-Unconstrained gridded population data at the cell-level?: A simulation analysis in urban Namibia. PLoS One 17, e0271504 (2022).',
        'Z. Bai, J. Wang, M. Wang, M. Gao, J. Sun, Accuracy assessment of multi-source Gridded Population distribution datasets in China. Sustainability 10, 1363 (2018).',
        'G. B. West, Scale: The Universal Laws of Growth, Innovation, Sustainability, and the Pace of Life in Organisms, Cities, Economies, and Companies (Penguin, 2017).',
        'G. Galilei, Dialogues Concerning Two New Sciences (Dover, 1914).',
        'K. J. Niklas, S. T. Hammond, On the interpretation of the normalization constant in the scaling equation. Front. Ecol. Evol. 6 (2019).',
        'D. S. Glazier, Commentary: On the Interpretation of the Normalization Constant in the Scaling Equation. Frontiers in Ecology and Evolution 8 (2020).',
        'J. A. Thomson, On Growth and Form. Nature 100, 21\u201322 (1917).',
        'T. J. Horder, "D\'arcy Wentworth Thompson, On growth and form, (1917)" in Landmark Writings in Western Mathematics 1640-1940 (Elsevier, 2005), pp. 823\u2013832.',
        'M. Kleiber, Body size and metabolism. Hilgardia (1932).',
        'J. Huxley, Problems of Relative Growth (Dover, 1972).',
        'L. M. A. Bettencourt, J. Lobo, D. Helbing, C. K\u00fchnert, G. B. West, Growth, innovation, scaling, and the pace of life in cities. Proc. Natl. Acad. Sci. U. S. A. 104, 7301\u20137306 (2007).',
        'L. M. A. Bettencourt, Introduction to Urban Science. [Preprint] (2021). https://doi.org/10.7551/mitpress/13909.001.0001.',
        'L. M. A. Bettencourt, The origins of scaling in cities. Science 340, 1438\u20131441 (2013).',
        'P. Kaitaniemi, How to derive biological information from the value of the normalization constant in allometric equations. PLoS One 3, e1932 (2008).',
        'H. Lumer, The Relation between b and k in Systems of Relative Growth Functions of the Form y = bx^k. Am. Nat. 70, 188\u2013191 (1936).',
        'G. K. Zipf, Human behavior and the principle of least effort: An introduction to human ecology. (1949).',
        # New references added for revision
        'L. M. A. Bettencourt, J. Lobo, Urban scaling in Europe. J. R. Soc. Interface 13, 20160005 (2016).',
        'R. Mao, Y. Bao, H. Duan, et al., Global urban subway development, construction material stocks, and embodied carbon emissions. Humanit. Soc. Sci. Commun. 8, 131 (2021).',
        'T. Fishman, A. Mastrucci, Y. Peled, G. Kamiya, M. Fridman, D. B. M\u00fcller, RASMI: Global ranges of building material intensities differentiated by region, structure, and function. Sci. Data 11, 418 (2024).',
        'A. Clauset, C. R. Shalizi, M. E. J. Newman, Power-law distributions in empirical data. SIAM Rev. 51, 661\u2013703 (2009).',
        'X. Gabaix, R. Ibragimov, Rank-1/2: A simple way to improve the OLS estimation of tail exponents. J. Bus. Econ. Stat. 29, 24\u201339 (2011).',
        'M. C. Gonz\u00e1lez, C. A. Hidalgo, A.-L. Barab\u00e1si, Understanding individual human mobility patterns. Nature 453, 779\u2013782 (2008).',
        'L. M. A. Bettencourt, V. C. Yang, J. Lobo, C. P. Kempes, D. Rybski, M. J. Hamilton, The interpretation of urban scaling analysis in time. J. R. Soc. Interface 17, 20190846 (2020).',
        'J. Lobo, L. Bettencourt, S. G. Ortman, Urban scaling theory: Answers to frequent questions. Environ. Plan. B 52, 1701\u20131716 (2024).',
    ]

    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph()
        r = p.add_run(f'{i}. ')
        r.bold = True
        r.font.size = Pt(10)
        r.font.name = 'Times New Roman'
        r2 = p.add_run(ref)
        r2.font.size = Pt(10)
        r2.font.name = 'Times New Roman'
        p.paragraph_format.space_after = Pt(3)

    # =========================================================================
    # SAVE
    # =========================================================================
    today = datetime.date.today().strftime('%Y-%m-%d')
    outpath = os.path.join(OUT_DIR, f'SI_Revised_Complete_{today}.docx')
    doc.save(outpath)
    print(f'Document saved to: {outpath}')
    print(f'Total sections: 10')
    print(f'New figures: S7-S15 (9 new + 6 existing = 15 total)')
    print(f'New tables: S4-S12 (9 new + 3 existing = 12 total)')
    return outpath


if __name__ == '__main__':
    build_si()
